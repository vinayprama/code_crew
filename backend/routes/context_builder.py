
import os
from typing import List, Dict

from utils.vector_store import search_similar_chunks
import boto3
from botocore.exceptions import BotoCoreError, ClientError
from io import BytesIO
from docx import Document
from dotenv import load_dotenv
load_dotenv()

AWS_REGION = os.getenv("AWS_REGION")
AWS_ACCESS_KEY_ID = os.getenv("AWS_ACCESS_KEY_ID")
AWS_SECRET_ACCESS_KEY = os.getenv("AWS_SECRET_ACCESS_KEY")
S3_BUCKET_NAME = os.getenv("S3_BUCKET_NAME")
GLOBAL_DOC_CONTEXT = "This is fallback content. No project documents were found."

s3_client = boto3.client(
    "s3",
    region_name=AWS_REGION,
    aws_access_key_id=AWS_ACCESS_KEY_ID,
    aws_secret_access_key=AWS_SECRET_ACCESS_KEY,
)

# def list_user_project_files(project_name: str) -> List[str]:
#     prefix = f"code_crew/{project_name}/"
#     try:
#         response = s3_client.list_objects_v2(Bucket=S3_BUCKET_NAME, Prefix=prefix)
#         return [item["Key"] for item in response.get("Contents", [])]
#     except (BotoCoreError, ClientError) as e:
#         print(f"⚠️ S3 list error: {e}")
#         return []

# def read_s3_file_content(key: str) -> str:
#     try:
#         response = s3_client.get_object(Bucket=S3_BUCKET_NAME, Key=key)
#         file_bytes = response["Body"].read()

#         if key.endswith(".docx"):
#             doc = Document(BytesIO(file_bytes))
#             doc_text = "\n".join([para.text for para in doc.paragraphs])
#             return doc_text

#         return file_bytes.decode("utf-8", errors="ignore")

#     except Exception as e:
#         print(f"❌ S3 read error {key}: {e}")
#         return ""
    

# async def build_context(project_name: str, query: str) -> str:
#     try:
#         file_keys = list_user_project_files(project_name)

#         if not file_keys:
#             print(f"⚠️ No files found for project '{project_name}' in S3.")
#             return GLOBAL_DOC_CONTEXT

#         all_text = ""
#         for key in file_keys:
#             file_text = read_s3_file_content(key)
#             all_text += f"\n{file_text}"

#         if not all_text.strip():
#             return GLOBAL_DOC_CONTEXT

#         # Now pass full document text to vector store for similarity search
#         similar_chunks: List[Dict] = search_similar_chunks(all_text, query)
        


#         if not similar_chunks:
#             return "(No relevant content found in vector store.)"

#         chunk_texts = [chunk.payload.get("text", "") for chunk in similar_chunks if chunk.payload and "text" in chunk.payload]
       

#         context = "\n\n".join(chunk_texts)
#         return context if context else "(No relevant content found.)"

#     except Exception as e:
#         print(f"Error in build_context: {str(e)}")
#         return f"Failed to build context: {str(e)}"

def list_user_project_files(project_name: str) -> List[str]:
    prefix = f"code_crew/{project_name.strip().replace(' ', '_')}/"
    try:
        response = s3_client.list_objects_v2(Bucket=S3_BUCKET_NAME, Prefix=prefix)
        return [item["Key"] for item in response.get("Contents", [])]
    except (BotoCoreError, ClientError) as e:
        print(f"⚠️ S3 list error: {e}")
        return []

def read_s3_file_content(key: str) -> str:
    try:
        response = s3_client.get_object(Bucket=S3_BUCKET_NAME, Key=key)
        file_bytes = response["Body"].read()

        if key.endswith(".docx"):
            doc = Document(BytesIO(file_bytes))
            return "\n".join([para.text for para in doc.paragraphs])

        if key.endswith(".txt"):
            return file_bytes.decode("utf-8", errors="ignore")

        return ""  # skip other types

    except Exception as e:
        print(f"❌ S3 read error {key}: {e}")
        return ""

async def build_context(project_name: str, query: str) -> str:
    try:
        file_keys = list_user_project_files(project_name)

        if not file_keys:
            print(f"⚠️ No files found for project '{project_name}' in S3.")
            return GLOBAL_DOC_CONTEXT

        all_text = ""
        for key in file_keys:
            if key.endswith((".docx", ".txt")):  # only read docx or transcript files
                file_text = read_s3_file_content(key)
                all_text += f"\n{file_text}"

        if not all_text.strip():
            return GLOBAL_DOC_CONTEXT

        similar_chunks: List[Dict] = search_similar_chunks(all_text, query)
        if not similar_chunks:
            return "(No relevant content found in vector store.)"

        chunk_texts = [chunk.payload.get("text", "") for chunk in similar_chunks if chunk.payload and "text" in chunk.payload]        
        context = "\n\n".join(chunk_texts)
        return context if context else "(No relevant content found.)"

    except Exception as e:
        print(f"Error in build_context: {str(e)}")
        return f"Failed to build context: {str(e)}"